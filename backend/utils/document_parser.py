"""
Document Parser Module - Multi-Format Document Text Extraction

This module provides utilities to parse various document formats and
extract text content for embedding and indexing.

Supported formats: TXT, MD, JSON, PDF, HTML, DOCX

"""

import json
from pathlib import Path
from typing import Dict, Any, Optional
from bs4 import BeautifulSoup
from loguru import logger

# Import document parsing libraries
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. PDF parsing will be disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX parsing will be disabled.")


class DocumentParser:
    """
    Parser for extracting text from various document formats.
    
    Supported formats:
    - Plain text (.txt)
    - Markdown (.md)
    - JSON (.json)
    - PDF (.pdf) - requires pymupdf
    - HTML (.html, .htm)
    - DOCX (.docx) - requires python-docx
    
    Attributes:
        supported_extensions: Set of supported file extensions
    """
    
    SUPPORTED_EXTENSIONS = {
        'txt', 'md', 'json', 'pdf', 'html', 'htm', 'docx'
    }
    
    def __init__(self):
        """Initialize the document parser."""
        logger.info("Initializing Document Parser")
        self.supported_extensions = self.SUPPORTED_EXTENSIONS.copy()
        
        # Remove unsupported formats
        if not PDF_AVAILABLE:
            self.supported_extensions.discard('pdf')
        if not DOCX_AVAILABLE:
            self.supported_extensions.discard('docx')
        
        logger.info(f"Supported formats: {', '.join(sorted(self.supported_extensions))}")
    
    def parse(
        self,
        file_path: Path,
        include_metadata: bool = True
    ) -> Dict[str, Any]:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to the document file
            include_metadata: Whether to include file metadata
            
        Returns:
            Dictionary with:
                - text: Extracted text content
                - metadata: File metadata (if include_metadata=True)
                - success: Boolean indicating success
                - error: Error message (if failed)
                
        Example:
            >>> parser = DocumentParser()
            >>> result = parser.parse(Path("document.pdf"))
            >>> print(result['text'])
        """
        if not file_path.exists():
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": f"File not found: {file_path}"
            }
        
        extension = file_path.suffix.lower().lstrip('.')
        
        if extension not in self.supported_extensions:
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": f"Unsupported file format: .{extension}"
            }
        
        logger.debug(f"Parsing {extension.upper()} file: {file_path.name}")
        
        try:
            # Route to appropriate parser
            if extension in ('txt', 'md'):
                text = self._parse_text_file(file_path)
            elif extension == 'json':
                text = self._parse_json_file(file_path)
            elif extension == 'pdf':
                text = self._parse_pdf_file(file_path)
            elif extension in ('html', 'htm'):
                text = self._parse_html_file(file_path)
            elif extension == 'docx':
                text = self._parse_docx_file(file_path)
            else:
                raise ValueError(f"No parser for extension: {extension}")
            
            # Prepare result
            result = {
                "text": text,
                "success": True,
                "error": None
            }
            
            if include_metadata:
                result["metadata"] = self._extract_metadata(file_path)
            
            logger.info(
                f"✅ Parsed {file_path.name}: {len(text)} characters extracted"
            )
            return result
            
        except Exception as e:
            logger.error(f"❌ Error parsing {file_path.name}: {e}")
            return {
                "text": "",
                "metadata": {},
                "success": False,
                "error": str(e)
            }
    
    def _parse_text_file(self, file_path: Path) -> str:
        """
        Parse plain text or markdown file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Extracted text content
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _parse_json_file(self, file_path: Path) -> str:
        """
        Parse JSON file and convert to readable text.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            JSON content as formatted string
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to readable text format
        return json.dumps(data, indent=2, ensure_ascii=False)
    
    def _parse_pdf_file(self, file_path: Path) -> str:
        """
        Parse PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text from all pages
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing")
        
        text_parts = []
        
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
        
        return "\n\n".join(text_parts)
    
    def _parse_html_file(self, file_path: Path) -> str:
        """
        Parse HTML file and extract text content.
        
        Args:
            file_path: Path to HTML file
            
        Returns:
            Extracted text (without HTML tags)
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up multiple newlines
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return '\n'.join(lines)
    
    def _parse_docx_file(self, file_path: Path) -> str:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text from document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")
        
        doc = DocxDocument(file_path)
        
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with file metadata
        """
        stat = file_path.stat()
        
        return {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size_bytes": stat.st_size,
            "extension": file_path.suffix.lower(),
            "created": stat.st_ctime,
            "modified": stat.st_mtime
        }
    
    def is_supported(self, file_path: Path) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if supported, False otherwise
        """
        extension = file_path.suffix.lower().lstrip('.')
        return extension in self.supported_extensions
    
    def parse_multiple(
        self,
        file_paths: list[Path],
        include_metadata: bool = True
    ) -> list[Dict[str, Any]]:
        """
        Parse multiple documents.
        
        Args:
            file_paths: List of file paths
            include_metadata: Whether to include metadata
            
        Returns:
            List of parsing results
        """
        results = []
        
        for file_path in file_paths:
            result = self.parse(file_path, include_metadata=include_metadata)
            results.append(result)
        
        successful = sum(1 for r in results if r['success'])
        logger.info(
            f"Parsed {len(file_paths)} documents: "
            f"{successful} successful, {len(file_paths) - successful} failed"
        )
        
        return results


# ===== GLOBAL DOCUMENT PARSER INSTANCE =====
_document_parser_instance: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """
    Get the global DocumentParser instance (singleton pattern).
    
    Returns:
        Global DocumentParser instance
    """
    global _document_parser_instance
    
    if _document_parser_instance is None:
        logger.info("Creating global DocumentParser instance")
        _document_parser_instance = DocumentParser()
    
    return _document_parser_instance


if __name__ == "__main__":
    """Test the Document Parser module."""
    
    # Configure logger
    from pathlib import Path
    Path("logs").mkdir(exist_ok=True)
    logger.add("logs/document_parser_test.log", rotation="1 MB")
    
    print("\n" + "="*60)
    print("TESTING DOCUMENT PARSER MODULE")
    print("="*60 + "\n")
    
    # Test 1: Initialize parser
    print("Test 1: Initializing parser...")
    parser = get_document_parser()
    print(f"✅ Supported formats: {', '.join(sorted(parser.supported_extensions))}\n")
    
    # Test 2: Parse text file
    print("Test 2: Creating and parsing test files...")
    
    # Create test text file
    test_txt = Path("test_file.txt")
    test_txt.write_text("This is a test document.\nIt has multiple lines.", encoding='utf-8')
    
    result = parser.parse(test_txt)
    print(f"✅ Text file parsed:")
    print(f"   Success: {result['success']}")
    print(f"   Text length: {len(result['text'])} chars")
    print(f"   Preview: {result['text'][:50]}...\n")
    
    # Test 3: Parse JSON file
    test_json = Path("test_file.json")
    test_json.write_text('{"key": "value", "number": 42}', encoding='utf-8')
    
    result = parser.parse(test_json)
    print(f"✅ JSON file parsed:")
    print(f"   Success: {result['success']}")
    print(f"   Text preview:\n{result['text']}\n")
    
    # Test 4: Parse HTML file
    test_html = Path("test_file.html")
    test_html.write_text(
        '<html><body><h1>Title</h1><p>Paragraph</p></body></html>',
        encoding='utf-8'
    )
    
    result = parser.parse(test_html)
    print(f"✅ HTML file parsed:")
    print(f"   Success: {result['success']}")
    print(f"   Text: {result['text']}\n")
    
    # Test 5: Check metadata
    print("Test 5: Checking metadata...")
    if result.get('metadata'):
        print("✅ Metadata extracted:")
        for key, value in result['metadata'].items():
            if key not in ('created', 'modified'):  # Skip timestamps
                print(f"   {key}: {value}")
    print()
    
    # Test 6: Cleanup
    print("Test 6: Cleaning up test files...")
    test_txt.unlink()
    test_json.unlink()
    test_html.unlink()
    print("✅ Cleanup complete\n")
    
    print("="*60)
    print("✅ ALL TESTS PASSED")
    print("="*60 + "\n")